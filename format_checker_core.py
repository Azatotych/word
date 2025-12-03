"""Ядро проверки оформления DOCX: функции анализа и аннотации."""
import math
import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import List, Optional, Tuple

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

EMU_PER_INCH = 914400
MM_PER_INCH = 25.4
DEFAULT_BODY_FONT = "Times New Roman"
EXPECTED_BODY_SIZE_PT = 10
EXPECTED_CAPTION_SIZE_PT = 9
EXPECTED_TITLE_SIZE_PT = 13
EXPECTED_INDENT_CM = 0.5
MARGIN_TOLERANCE_CM = 0.2
PAGE_TOLERANCE_MM = 1.0
LINES_PER_PAGE = 35
LINE_LENGTH_CHARS = 70


@dataclass
class Issue:
    """Структура замечания по оформлению."""

    rule: str
    level: str
    message: str
    paragraph_index: Optional[int] = None
    paragraph_text: Optional[str] = None

    def to_dict(self, file_name: Optional[str] = None) -> dict:
        data = asdict(self)
        if file_name is not None:
            data["file"] = file_name
        return {key: value for key, value in data.items() if value is not None}


def emu_to_mm(emu: int) -> float:
    return emu / EMU_PER_INCH * MM_PER_INCH


def approx_equal(value: float, expected: float, tolerance: float) -> bool:
    return abs(value - expected) <= tolerance


def load_document(path: str) -> docx.document.Document:
    """Загружает документ DOCX."""

    return docx.Document(path)


def estimate_page_count(doc: docx.document.Document) -> Tuple[int, int]:
    """Оценивает количество страниц и число абзацев на последней странице."""

    total_lines = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        lines = max(1, math.ceil(len(text) / LINE_LENGTH_CHARS))
        total_lines += lines

    page_count = max(1, math.ceil(total_lines / LINES_PER_PAGE))

    threshold = (page_count - 1) * LINES_PER_PAGE
    accumulated = 0
    last_page_paragraphs = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        lines = max(1, math.ceil(len(text) / LINE_LENGTH_CHARS))
        accumulated += lines
        if accumulated > threshold:
            last_page_paragraphs += 1

    return page_count, last_page_paragraphs


def get_effective_font(paragraph) -> Tuple[Optional[str], Optional[float], Optional[bool], Optional[bool]]:
    """Возвращает оценку шрифта (имя, размер в pt, bold, italic) для абзаца."""

    names = []
    sizes = []
    bold_flags = []
    italic_flags = []
    for run in paragraph.runs:
        if not run.text:
            continue
        if run.font.name:
            names.append(run.font.name)
        if run.font.size:
            sizes.append(run.font.size.pt)
        if run.font.bold is not None:
            bold_flags.append(run.font.bold)
        if run.font.italic is not None:
            italic_flags.append(run.font.italic)

    if not names and paragraph.style and paragraph.style.font and paragraph.style.font.name:
        names.append(paragraph.style.font.name)
    if not sizes and paragraph.style and paragraph.style.font and paragraph.style.font.size:
        sizes.append(paragraph.style.font.size.pt)
    if not bold_flags and paragraph.style and paragraph.style.font and paragraph.style.font.bold is not None:
        bold_flags.append(paragraph.style.font.bold)
    if not italic_flags and paragraph.style and paragraph.style.font and paragraph.style.font.italic is not None:
        italic_flags.append(paragraph.style.font.italic)

    name = names[0] if names else None
    size = sizes[0] if sizes else None
    bold = bold_flags[0] if bold_flags else None
    italic = italic_flags[0] if italic_flags else None
    return name, size, bold, italic


def check_page_setup(doc: docx.document.Document) -> List[Issue]:
    issues: List[Issue] = []
    page_size_ok = True
    margins_ok = True

    for idx, section in enumerate(doc.sections, start=1):
        width_mm = emu_to_mm(section.page_width)
        height_mm = emu_to_mm(section.page_height)
        if not (approx_equal(width_mm, 148, PAGE_TOLERANCE_MM) and approx_equal(height_mm, 210, PAGE_TOLERANCE_MM)):
            page_size_ok = False
            issues.append(
                Issue(
                    rule="PAGE_SIZE",
                    level="ERROR",
                    message=f"Секция {idx}: размер страницы {width_mm:.1f}×{height_mm:.1f} мм, ожидается A5 148×210 мм",
                )
            )

        top_cm = section.top_margin.cm if section.top_margin is not None else 0
        bottom_cm = section.bottom_margin.cm if section.bottom_margin is not None else 0
        left_cm = section.left_margin.cm if section.left_margin is not None else 0
        right_cm = section.right_margin.cm if section.right_margin is not None else 0
        if not (
            approx_equal(top_cm, 1.6, MARGIN_TOLERANCE_CM)
            and approx_equal(bottom_cm, 1.4, MARGIN_TOLERANCE_CM)
            and approx_equal(left_cm, 1.5, MARGIN_TOLERANCE_CM)
            and approx_equal(right_cm, 1.5, MARGIN_TOLERANCE_CM)
        ):
            margins_ok = False
            issues.append(
                Issue(
                    rule="MARGINS",
                    level="ERROR",
                    message=(
                        f"Секция {idx}: поля верх/низ/лево/право = "
                        f"{top_cm:.2f}/{bottom_cm:.2f}/{left_cm:.2f}/{right_cm:.2f} см, ожидается 1.6/1.4/1.5/1.5 см"
                    ),
                )
            )

    if page_size_ok:
        issues.append(Issue(rule="PAGE_SIZE", level="OK", message="Размер страницы соответствует формату A5"))
    if margins_ok:
        issues.append(Issue(rule="MARGINS", level="OK", message="Поля страницы соответствуют требованиям"))

    page_count, last_page_paragraphs = estimate_page_count(doc)
    if page_count > 5:
        issues.append(
            Issue(
                rule="PAGE_COUNT",
                level="ERROR",
                message=f"Оценочное количество страниц: {page_count}, допускается не более 5",
            )
        )
    else:
        issues.append(Issue(rule="PAGE_COUNT", level="OK", message=f"Оценочное количество страниц: {page_count}"))

    if page_count > 1 and last_page_paragraphs < 3:
        issues.append(
            Issue(
                rule="LAST_PAGE_FILL",
                level="WARN",
                message="Последняя страница содержит очень мало текста (менее 3 абзацев)",
            )
        )

    return issues


def detect_structure_indices(doc: docx.document.Document) -> Tuple[Optional[int], Optional[int], Optional[int]]:
    """Возвращает индексы абзацев: авторы, заголовок, первый абзац тела."""

    paragraphs = doc.paragraphs
    first_nonempty = next((i for i, p in enumerate(paragraphs) if p.text.strip()), None)
    if first_nonempty is None:
        return None, None, None

    authors_idx = first_nonempty
    title_idx = None
    body_idx = None

    idx = authors_idx + 1
    while idx < len(paragraphs) and not paragraphs[idx].text.strip():
        idx += 1

    if idx < len(paragraphs):
        title_idx = idx
        idx += 1
        while idx < len(paragraphs) and not paragraphs[idx].text.strip():
            idx += 1
        if idx < len(paragraphs):
            body_idx = idx

    return authors_idx, title_idx, body_idx


def check_paragraphs(doc: docx.document.Document) -> List[Issue]:
    issues: List[Issue] = []
    authors_idx, title_idx, _ = detect_structure_indices(doc)

    literature_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().lower() == "литература":
            literature_idx = i
            break

    font_body_ok = True
    indent_ok = True
    spacing_ok = True
    font9_warned = False

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        is_figure_caption = bool(re.match(r"^(рис\.|рисунок)\s", text, re.IGNORECASE))
        in_literature = literature_idx is not None and i > literature_idx
        is_literature_header = literature_idx == i
        is_special = i in {authors_idx, title_idx} or is_figure_caption or is_literature_header or in_literature

        name, size, bold, italic = get_effective_font(paragraph)

        if not is_special:
            if size is not None and not approx_equal(size, EXPECTED_BODY_SIZE_PT, 0.5):
                font_body_ok = False
                issues.append(
                    Issue(
                        rule="FONT_BODY",
                        level="ERROR",
                        message=f"Абзац {i + 1}: размер шрифта {size:.1f} pt, ожидается 10 pt",
                        paragraph_index=i,
                        paragraph_text=text[:80],
                    )
                )
            if name is not None and DEFAULT_BODY_FONT.lower() not in name.lower():
                font_body_ok = False
                issues.append(
                    Issue(
                        rule="FONT_BODY",
                        level="ERROR",
                        message=f"Абзац {i + 1}: шрифт '{name}', ожидается Times New Roman",
                        paragraph_index=i,
                        paragraph_text=text[:80],
                    )
                )

        if in_literature or is_literature_header or is_figure_caption:
            expected_size = EXPECTED_CAPTION_SIZE_PT
            if size is not None and not approx_equal(size, expected_size, 0.5):
                font9_warned = True
                issues.append(
                    Issue(
                        rule="FONT_9PT",
                        level="WARN",
                        message=f"Абзац {i + 1}: размер шрифта {size:.1f} pt, ожидается 9 pt",
                        paragraph_index=i,
                        paragraph_text=text[:80],
                    )
                )

        line_spacing = paragraph.paragraph_format.line_spacing
        if line_spacing not in (None, 1):
            if not (isinstance(line_spacing, (int, float)) and approx_equal(float(line_spacing), 1.0, 0.05)):
                spacing_ok = False
                issues.append(
                    Issue(
                        rule="LINE_SPACING",
                        level="WARN",
                        message=f"Абзац {i + 1}: межстрочный интервал задан как {line_spacing}",
                        paragraph_index=i,
                        paragraph_text=text[:80],
                    )
                )

        indent = paragraph.paragraph_format.first_line_indent.cm if paragraph.paragraph_format.first_line_indent else 0
        if not is_special:
            if indent == 0:
                if text.startswith("\t") or re.match(r"^ {3,}", paragraph.text):
                    indent_ok = False
                    issues.append(
                        Issue(
                            rule="INDENT_BY_TABS_OR_SPACES",
                            level="ERROR",
                            message=f"Абзац {i + 1}: отступ первой строки сделан табуляцией/пробелами",
                            paragraph_index=i,
                            paragraph_text=text[:80],
                        )
                    )
            elif not approx_equal(indent, EXPECTED_INDENT_CM, 0.05):
                indent_ok = False
                issues.append(
                    Issue(
                        rule="INDENT_SIZE",
                        level="ERROR",
                        message=f"Абзац {i + 1}: отступ первой строки {indent:.2f} см, ожидается 0.50 см",
                        paragraph_index=i,
                        paragraph_text=text[:80],
                    )
                )

    if font_body_ok:
        issues.append(Issue(rule="FONT_BODY", level="OK", message="Основной текст использует Times New Roman 10 pt"))
    if not font9_warned:
        issues.append(Issue(rule="FONT_9PT", level="OK", message="Элементы со шрифтом 9 pt не выявили отклонений"))
    if spacing_ok:
        issues.append(Issue(rule="LINE_SPACING", level="OK", message="Межстрочный интервал не отличается от одинарного"))
    if indent_ok:
        issues.append(Issue(rule="INDENT_SIZE", level="OK", message="Отступ первой строки соответствует требованию"))

    return issues


def check_structure(doc: docx.document.Document) -> List[Issue]:
    issues: List[Issue] = []
    authors_idx, title_idx, body_idx = detect_structure_indices(doc)
    paragraphs = doc.paragraphs

    if authors_idx is None:
        issues.append(Issue(rule="AUTHORS_LINE", level="WARN", message="Не найден первый непустой абзац с авторами"))
        return issues

    authors_par = paragraphs[authors_idx]
    name, size, bold, italic = get_effective_font(authors_par)
    author_text = authors_par.text.strip()
    author_ok = True

    if authors_par.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
        author_ok = False
    if size is not None and not approx_equal(size, EXPECTED_BODY_SIZE_PT, 0.5):
        author_ok = False
    if not (bold and italic):
        author_ok = False
    if not re.search(r"[А-ЯЁ][а-яё]+", author_text):
        author_ok = False

    if not author_ok:
        issues.append(
            Issue(
                rule="AUTHORS_LINE",
                level="WARN",
                message="Строка с фамилиями авторов не соответствует требованиям (выравнивание/шрифт/формат)",
                paragraph_index=authors_idx,
                paragraph_text=author_text[:80],
            )
        )
    else:
        issues.append(Issue(rule="AUTHORS_LINE", level="OK", message="Строка авторов оформлена корректно"))

    if title_idx is None:
        issues.append(Issue(rule="TITLE_LAYOUT", level="ERROR", message="Не удалось найти заголовок статьи"))
        return issues

    has_blank_between = any(not p.text.strip() for p in paragraphs[authors_idx + 1 : title_idx])
    if not has_blank_between:
        issues.append(Issue(rule="TITLE_LAYOUT", level="WARN", message="После строки авторов отсутствует пустая строка"))

    title_par = paragraphs[title_idx]
    title_name, title_size, title_bold, title_italic = get_effective_font(title_par)
    title_text = title_par.text.strip()
    title_ok = True
    if title_par.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        title_ok = False
    if title_size is not None and not approx_equal(title_size, EXPECTED_TITLE_SIZE_PT, 0.5):
        title_ok = False
    if title_bold or title_italic:
        title_ok = False
    if "-\n" in title_text:
        issues.append(Issue(rule="TITLE_LAYOUT", level="WARN", message="В заголовке встречаются ручные переносы"))

    if not title_ok:
        issues.append(
            Issue(
                rule="TITLE_FORMAT",
                level="ERROR",
                message="Заголовок не соответствует требуемому шрифту/выравниванию",
                paragraph_index=title_idx,
                paragraph_text=title_text[:80],
            )
        )
    else:
        issues.append(Issue(rule="TITLE_FORMAT", level="OK", message="Заголовок оформлен корректно"))

    has_blank_after_title = False
    if title_idx + 1 < len(paragraphs):
        has_blank_after_title = paragraphs[title_idx + 1].text.strip() == ""
    if not has_blank_after_title:
        issues.append(Issue(rule="TITLE_SPACING", level="WARN", message="После заголовка нет пустой строки"))
    else:
        issues.append(Issue(rule="TITLE_SPACING", level="OK", message="Отступ после заголовка присутствует"))

    if body_idx is not None:
        body_par = paragraphs[body_idx]
        body_name, body_size, _, _ = get_effective_font(body_par)
        body_text = body_par.text.strip()
        body_ok = True
        if body_par.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            body_ok = False
        if body_size is not None and not approx_equal(body_size, EXPECTED_BODY_SIZE_PT, 0.5):
            body_ok = False
        if body_name is not None and DEFAULT_BODY_FONT.lower() not in body_name.lower():
            body_ok = False

        if not body_ok:
            issues.append(
                Issue(
                    rule="BODY_START",
                    level="ERROR",
                    message="Первый абзац основного текста не соответствует требованиям",
                    paragraph_index=body_idx,
                    paragraph_text=body_text[:80],
                )
            )
        else:
            issues.append(Issue(rule="BODY_START", level="OK", message="Начало основного текста оформлено корректно"))
    else:
        issues.append(Issue(rule="BODY_START", level="ERROR", message="Не удалось определить начало основного текста"))

    return issues


def check_literature(doc: docx.document.Document) -> List[Issue]:
    issues: List[Issue] = []
    paragraphs = doc.paragraphs
    literature_idx = None

    for i, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "Литература":
            literature_idx = i
            name, size, bold, _ = get_effective_font(paragraph)
            header_ok = True
            if size is not None and not approx_equal(size, EXPECTED_CAPTION_SIZE_PT, 0.5):
                header_ok = False
            if not bold:
                header_ok = False
            if not header_ok:
                issues.append(
                    Issue(
                        rule="LITERATURE_HEADER",
                        level="ERROR",
                        message="Заголовок 'Литература' оформлен неверно (требуется 9 pt и полужирный)",
                        paragraph_index=i,
                        paragraph_text=paragraph.text.strip()[:80],
                    )
                )
            else:
                issues.append(Issue(rule="LITERATURE_HEADER", level="OK", message="Заголовок 'Литература' оформлен корректно"))
            break

    if literature_idx is None:
        issues.append(Issue(rule="LITERATURE_HEADER", level="ERROR", message="Не найден заголовок 'Литература'"))
        return issues

    items_ok = True
    for idx, paragraph in enumerate(paragraphs[literature_idx + 1 :]):
        text = paragraph.text.strip()
        if not text:
            continue
        name, size, _, _ = get_effective_font(paragraph)
        if size is not None and not approx_equal(size, EXPECTED_CAPTION_SIZE_PT, 0.5):
            items_ok = False
        if paragraph.alignment not in (WD_ALIGN_PARAGRAPH.JUSTIFY, None):
            items_ok = False
        if not items_ok:
            issues.append(
                Issue(
                    rule="LITERATURE_ITEMS",
                    level="WARN",
                    message="Элементы списка литературы должны быть 9 pt и выровнены по ширине",
                    paragraph_index=literature_idx + 1 + idx,
                    paragraph_text=text[:80],
                )
            )
            break

    if items_ok:
        issues.append(Issue(rule="LITERATURE_ITEMS", level="OK", message="Список литературы соответствует требованиям"))

    return issues


def check_figures(doc: docx.document.Document) -> List[Issue]:
    issues: List[Issue] = []
    found_caption = False
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if re.match(r"^(рис\.|рисунок)\s", text, re.IGNORECASE):
            found_caption = True
            name, size, _, _ = get_effective_font(paragraph)
            if size is not None and not approx_equal(size, EXPECTED_CAPTION_SIZE_PT, 0.5):
                issues.append(
                    Issue(
                        rule="FIGURE_CAPTION_FONT",
                        level="WARN",
                        message=f"Подпись к рисунку (абзац {i + 1}) должна быть 9 pt",
                        paragraph_index=i,
                        paragraph_text=text[:80],
                    )
                )
            if text.rstrip().endswith('.'):
                issues.append(
                    Issue(
                        rule="FIGURE_CAPTION_DOT",
                        level="WARN",
                        message=f"Подпись к рисунку (абзац {i + 1}) заканчивается точкой",
                        paragraph_index=i,
                        paragraph_text=text[:80],
                    )
                )
    if found_caption and not any(issue.rule.startswith("FIGURE_CAPTION") for issue in issues):
        issues.append(Issue(rule="FIGURE_CAPTION_FONT", level="OK", message="Подписи к рисункам оформлены корректно"))
    return issues


def check_special_text_rules(doc: docx.document.Document) -> List[Issue]:
    issues: List[Issue] = []
    tabs_found = False
    leading_spaces_found = False
    nbsp_warned = False
    dash_warned = False

    combined_text = "\n".join(p.text for p in doc.paragraphs)

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if "\t" in text and not tabs_found:
            tabs_found = True
            issues.append(
                Issue(
                    rule="TABS_IN_TEXT",
                    level="WARN",
                    message="Обнаружены символы табуляции в тексте",
                    paragraph_index=idx,
                    paragraph_text=text.strip()[:80],
                )
            )
        if re.match(r"^ {3,}", text) and not leading_spaces_found:
            leading_spaces_found = True
            issues.append(
                Issue(
                    rule="LEADING_SPACES",
                    level="WARN",
                    message="Абзацы начинаются с нескольких пробелов",
                    paragraph_index=idx,
                    paragraph_text=text.strip()[:80],
                )
            )

    pattern = re.compile(r"\b\d+(\s|\xa0)+(кг|г|мм|см|м|км|№|§)\b", re.IGNORECASE)
    for match in pattern.finditer(combined_text):
        if "\xa0" not in match.group(0):
            nbsp_warned = True
            issues.append(
                Issue(
                    rule="NONBREAKING_SPACE",
                    level="WARN",
                    message="Числа с единицами измерения оформлены обычным пробелом, требуется неразрывный",
                )
            )
            break

    if re.search(r"\s-\s", combined_text):
        dash_warned = True
        issues.append(Issue(rule="HYPHEN_DASH_MIX", level="WARN", message="Найдено возможное использование дефиса вместо тире"))

    if not tabs_found:
        issues.append(Issue(rule="TABS_IN_TEXT", level="OK", message="Табуляции в тексте не обнаружены"))
    if not leading_spaces_found:
        issues.append(Issue(rule="LEADING_SPACES", level="OK", message="Лишние пробелы в начале абзацев не обнаружены"))
    if not nbsp_warned:
        issues.append(Issue(rule="NONBREAKING_SPACE", level="OK", message="Неразрывные пробелы не вызывают нареканий"))
    if not dash_warned:
        issues.append(Issue(rule="HYPHEN_DASH_MIX", level="OK", message="Подозрительных тире не найдено"))

    # TODO: автоматическая проверка формул не реализована
    return issues


def collect_issues(doc: docx.document.Document) -> List[Issue]:
    issues: List[Issue] = []
    issues.extend(check_page_setup(doc))
    issues.extend(check_paragraphs(doc))
    issues.extend(check_structure(doc))
    issues.extend(check_literature(doc))
    issues.extend(check_figures(doc))
    issues.extend(check_special_text_rules(doc))
    return issues


def check_document(path: str) -> List[Issue]:
    """Выполняет полную проверку документа и возвращает список Issue."""

    doc = load_document(path)
    return collect_issues(doc)


def annotate_document(source_path: str, issues: List[Issue]) -> str:
    """Создаёт копию DOCX с подсветкой проблемных абзацев красным цветом."""

    output_path = str(Path(source_path).with_name(f"{Path(source_path).stem}_annotated{Path(source_path).suffix}"))
    doc = docx.Document(source_path)
    problem_paragraphs = {
        issue.paragraph_index
        for issue in issues
        if issue.paragraph_index is not None and issue.level != "OK"
    }

    for idx in problem_paragraphs:
        if idx is None or idx >= len(doc.paragraphs):
            continue
        paragraph = doc.paragraphs[idx]
        if not paragraph.runs:
            run = paragraph.add_run(paragraph.text)
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            continue
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.save(output_path)
    return output_path


def format_report(issues: List[Issue], file_name: str) -> Tuple[dict, str]:
    """Строит агрегированные данные отчёта для переиспользования CLI/GUI."""

    totals = {"ERROR": 0, "WARN": 0, "OK": 0}
    for issue in issues:
        totals[issue.level] = totals.get(issue.level, 0) + 1

    lines = [f"==== {file_name} ===="]
    if not issues:
        lines.append("Нарушений не обнаружено.\n")
    else:
        lines.append(
            "Итоги: ошибок = {errors}, предупреждений = {warns}, без замечаний = {oks}".format(
                errors=totals.get("ERROR", 0), warns=totals.get("WARN", 0), oks=totals.get("OK", 0)
            )
        )
        for issue in issues:
            location = f"(абзац {issue.paragraph_index + 1}) " if issue.paragraph_index is not None else ""
            snippet = f" | {issue.paragraph_text}" if issue.paragraph_text else ""
            lines.append(f"[{issue.level}] {issue.rule}: {location}{issue.message}{snippet}")
        lines.append("")

    return totals, "\n".join(lines)


__all__ = [
    "Issue",
    "annotate_document",
    "check_document",
    "format_report",
]
